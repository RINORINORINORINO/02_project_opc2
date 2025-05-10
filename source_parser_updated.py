import requests
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
import os
import docx
from typing import Union, Dict, List, Any, Optional, Tuple, Set
import time
import random
from urllib.parse import urlparse, urljoin
import re
from functools import lru_cache
import concurrent.futures
import logging
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry
from cloud_ocr import parse_cloud_ocr

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# 유튜브 처리를 위한 모듈 - 동적 임포트 처리
youtube_parser = None

# 타입 정의
SourceType = Union[str, Dict[str, str]]  # URL 문자열 또는 {'type': 'pdf', 'path': '...'} 형태

# 이 위치에 추가
SUPPORTED_FILE_TYPES = ['.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']

# 세션 객체 생성 및 재시도 설정
def create_session() -> requests.Session:
    """향상된 재시도 로직이 있는 세션 생성"""
    session = requests.Session()
    retry_strategy = Retry(
        total=3,  # 총 시도 횟수
        backoff_factor=0.5,  # 재시도 간격 증가 인자
        status_forcelist=[429, 500, 502, 503, 504],  # 재시도할 HTTP 상태 코드
        allowed_methods=["GET", "POST"]  # 재시도할 HTTP 메서드
    )
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    return session

# 기본 세션 생성
session = create_session()

@lru_cache(maxsize=32)
def parse_url(url: str) -> str:
    """
    웹 URL로부터 기사/본문 텍스트 추출 - 캐싱 및 오류 처리 강화
    
    Args:
        url: 파싱할 웹 URL
        
    Returns:
        추출된 텍스트 컨텐츠
    """
    # YouTube URL 확인
    if "youtube.com" in url or "youtu.be" in url:
        logger.info(f"🎬 YouTube 영상 URL 감지: {url}")
        return parse_youtube_content(url)
    
    try:
        logger.info(f"🌐 URL 파싱 시작: {url[:60]}{'...' if len(url) > 60 else ''}")
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Referer": "https://www.google.com/",
            "DNT": "1",
        }
        
        res = session.get(url, headers=headers, timeout=15)
        res.raise_for_status()  # 오류 상태 코드 확인
        
        # 인코딩 처리 (명시적 인코딩이 없는 경우 대비)
        if res.encoding.lower() == 'iso-8859-1':
            # 인코딩 감지 시도
            res.encoding = res.apparent_encoding
        
        soup = BeautifulSoup(res.text, "html.parser")
        
        # 미디어 플랫폼별 최적화된 파싱
        domain = urlparse(url).netloc
        
        # 특정 뉴스 사이트 최적화
        if "medium.com" in domain:
            content = parse_medium(soup)
        elif any(x in domain for x in ["nytimes.com", "washingtonpost.com", "theguardian.com"]):
            content = parse_news_site(soup)
        elif "wikipedia.org" in domain:
            content = parse_wikipedia(soup)
        elif "arxiv.org" in domain:
            content = parse_arxiv(soup)
        else:
            # 일반적인 파싱 방법
            content = general_parsing(soup, url)
            
        if not content:
            logger.warning(f"⚠️ {url}에서 콘텐츠를 찾을 수 없습니다. 일반 파싱으로 시도합니다.")
            content = general_parsing(soup, url)
            
        # 텍스트 정리
        content = clean_text(content)
        
        logger.info(f"✅ URL 파싱 완료: {url[:60]}{'...' if len(url) > 60 else ''}")
        return content

    except requests.exceptions.RequestException as e:
        logger.error(f"❌ URL 요청 오류 ({url}): {str(e)}")
        return f"URL 접근 오류: {url}\n오류 세부사항: {str(e)}"
    except Exception as e:
        logger.error(f"❌ URL 파싱 오류 ({url}): {str(e)}")
        return f"URL 파싱 오류: {url}\n오류 세부사항: {str(e)}"

def parse_medium(soup: BeautifulSoup) -> str:
    """Medium 아티클 파싱"""
    article = soup.select_one("article")
    if article:
        # Medium 제목 추출
        title_elem = soup.select_one("h1")
        title = title_elem.get_text(strip=True) if title_elem else ""
        
        # 본문 추출
        paragraphs = article.select("p")
        content = "\n".join(p.get_text(strip=True) for p in paragraphs)
        
        if title:
            return f"{title}\n\n{content}"
        return content
    return ""

def parse_news_site(soup: BeautifulSoup) -> str:
    """주요 뉴스 사이트 파싱"""
    # 기사 제목
    title_elem = soup.select_one("h1") or soup.select_one(".headline") or soup.select_one(".title")
    title = title_elem.get_text(strip=True) if title_elem else ""
    
    # 기사 본문 컨테이너 찾기
    article = soup.select_one("article") or soup.select_one(".article-body") or soup.select_one(".story-body")
    
    if article:
        # 불필요한 요소 제거
        for el in article.select(".ad, .advertisement, .social-share, .newsletter"):
            if el:
                el.decompose()
        
        paragraphs = article.select("p")
        content = "\n".join(p.get_text(strip=True) for p in paragraphs)
        
        if title:
            return f"{title}\n\n{content}"
        return content
    return ""

def parse_wikipedia(soup: BeautifulSoup) -> str:
    """위키피디아 파싱"""
    # 제목 추출
    title_elem = soup.select_one("#firstHeading")
    title = title_elem.get_text(strip=True) if title_elem else ""
    
    content_div = soup.select_one("#mw-content-text")
    if content_div:
        # 불필요한 요소 제거
        for el in content_div.select(".reference, .mw-editsection, .thumb, .navbox, .vertical-navbox"):
            el.decompose()
        
        # 본문 문단 추출
        paragraphs = content_div.select("p")
        content = "\n".join(p.get_text(strip=True) for p in paragraphs)
        
        # 목차 추출
        toc = []
        for heading in content_div.select("h2, h3"):
            if heading.get_text(strip=True) not in ["Contents", "References", "External links"]:
                toc.append(heading.get_text(strip=True).replace("[edit]", ""))
        
        toc_text = "목차: " + ", ".join(toc) if toc else ""
        
        result = []
        if title:
            result.append(title)
        if toc_text:
            result.append(toc_text)
        if content:
            result.append(content)
            
        return "\n\n".join(result)
    return ""

def parse_arxiv(soup: BeautifulSoup) -> str:
    """arXiv 논문 파싱"""
    # 논문 제목
    title_elem = soup.select_one(".title")
    title = title_elem.get_text(strip=True).replace("Title:", "").strip() if title_elem else ""
    
    # 저자
    authors_elem = soup.select_one(".authors")
    authors = authors_elem.get_text(strip=True).replace("Authors:", "").strip() if authors_elem else ""
    
    # 초록
    abstract = soup.select_one(".abstract")
    abstract_text = ""
    if abstract:
        abstract_text = abstract.get_text(strip=True).replace("Abstract: ", "")
    
    # 조합
    result = []
    if title:
        result.append(f"제목: {title}")
    if authors:
        result.append(f"저자: {authors}")
    if abstract_text:
        result.append(f"초록:\n{abstract_text}")
        
    return "\n\n".join(result)

def general_parsing(soup: BeautifulSoup, url: str) -> str:
    """일반적인 웹페이지 파싱 방법 - 도메인별 휴리스틱 개선"""
    # 제목 추출
    title = ""
    title_tag = soup.find("title")
    if title_tag:
        title = title_tag.get_text(strip=True)
    
    # meta description 추출 (페이지 요약에 유용)
    meta_desc = ""
    meta_tag = soup.find("meta", attrs={"name": "description"}) or soup.find("meta", attrs={"property": "og:description"})
    if meta_tag and meta_tag.get("content"):
        meta_desc = meta_tag["content"]
    
    # 도메인 정보 추출 (사이트 특성 파악)
    domain = urlparse(url).netloc
    
    # 도메인별 커스텀 컨텐츠 선택자
    domain_selectors = {
        "bbc.com": "article .ssrcss-uf6wea-RichTextComponentWrapper",
        "cnn.com": ".article__content",
        "reuters.com": ".article-body__content__17Yit",
        "bloomberg.com": ".body-content",
        "apnews.com": ".Article",
        "ft.com": ".article__content-body",
    }
    
    # 도메인별 맞춤 선택자 사용
    for site, selector in domain_selectors.items():
        if site in domain:
            content_elems = soup.select(selector)
            if content_elems:
                return title + "\n\n" + "\n".join(elem.get_text(strip=True) for elem in content_elems)
    
    # 여러 내용 컨테이너 후보 시도 (우선순위 순)
    content_candidates = [
        soup.select("article p"),
        soup.select("main p"),
        soup.select(".content p, .post-content p, .entry-content p, .article p"),
        soup.select(".story p, .body p, .post p"),
        soup.select('[role="main"] p'),
        soup.select('[role="article"] p'),
        soup.select("p")  # 마지막 수단
    ]
    
    for candidate in content_candidates:
        if candidate and len("".join(p.get_text() for p in candidate)) > 200:
            content = "\n".join(p.get_text(strip=True) for p in candidate)
            if title:
                return title + "\n\n" + content
            return content
    
    # 마지막 수단: 텍스트 덩어리 추출
    try:
        # 본문으로 추정되는 영역 식별
        main_content = identify_main_content(soup)
        
        if main_content:
            # 텍스트 추출 및 정리
            text = main_content.get_text(separator="\n", strip=True)
            lines = [line.strip() for line in text.split("\n") if len(line.strip()) > 30]
            
            if lines:
                if title:
                    return title + "\n\n" + "\n".join(lines)
                return "\n".join(lines)
    except Exception:
        pass
    
    # 모든 방법 실패 시 메타 정보라도 반환
    if title or meta_desc:
        result = []
        if title:
            result.append(f"제목: {title}")
        if meta_desc:
            result.append(f"설명: {meta_desc}")
        result.append(f"URL: {url}")
        result.append("(웹페이지에서 충분한 텍스트 콘텐츠를 추출하지 못했습니다.)")
        return "\n\n".join(result)
    
    return f"웹페이지 텍스트 추출 실패: {url}"

def identify_main_content(soup: BeautifulSoup) -> Optional[BeautifulSoup]:
    """
    본문 콘텐츠로 추정되는 영역을 식별하는 함수
    텍스트 길이와 밀도를 기준으로 판단
    """
    # 주요 콘텐츠 후보 태그
    candidates = {}
    
    # 일반적인 콘텐츠 컨테이너 태그
    for tag in ["article", "main", "div", "section"]:
        for element in soup.find_all(tag):
            # 클래스/ID 분석하여 광고, 네비게이션 등 제외
            if element.get("class"):
                class_str = " ".join(element.get("class")).lower()
                if any(x in class_str for x in ["nav", "menu", "sidebar", "footer", "comment", "ad-", "advertisement"]):
                    continue
            
            # id 분석
            if element.get("id"):
                id_str = element.get("id").lower()
                if any(x in id_str for x in ["nav", "menu", "sidebar", "footer", "comment"]):
                    continue
            
            # 텍스트 길이 계산
            text = element.get_text(" ", strip=True)
            if len(text) < 100:  # 너무 짧은 콘텐츠 제외
                continue
                
            # 텍스트 밀도 계산 (텍스트 길이 / HTML 길이)
            html_len = len(str(element))
            if html_len == 0:
                continue
                
            text_density = len(text) / html_len
            
            # 점수 계산 (텍스트 길이 * 텍스트 밀도)
            score = len(text) * text_density
            candidates[element] = score
    
    # 점수 기준 정렬
    if candidates:
        return max(candidates.items(), key=lambda x: x[1])[0]
    
    return None

def parse_pdf(path: str) -> str:
    """PDF 파일 경로에서 전체 텍스트 추출 - 향상된 버전"""
    try:
        logger.info(f"📄 PDF 파싱 시작: {os.path.basename(path)}")
        doc = fitz.open(path)
        texts = []
        
        # 메타데이터 추출
        metadata = doc.metadata
        if metadata.get("title"):
            texts.append(f"제목: {metadata.get('title')}")
        if metadata.get("author"):
            texts.append(f"저자: {metadata.get('author')}")
        if metadata.get("subject"):
            texts.append(f"주제: {metadata.get('subject')}")
        texts.append("")  # 빈 줄 추가
        
        # TOC(목차) 추출 시도
        toc = doc.get_toc()
        if toc:
            texts.append("목차:")
            for level, title, page in toc:
                indent = "  " * (level - 1)
                texts.append(f"{indent}- {title} (페이지: {page})")
            texts.append("")  # 빈 줄 추가
        
        # 본문 추출 - 텍스트 블록 기반 접근
        for page_num, page in enumerate(doc):
            # 페이지 구조 분석을 통한 향상된 텍스트 추출
            blocks = page.get_text("blocks")
            page_text = []
            
            for block in blocks:
                if block[6] == 0:  # 텍스트 블록
                    block_text = block[4].strip()
                    if block_text:
                        page_text.append(block_text)
            
            if page_text:
                # 페이지 번호 표시 (특히 긴 문서에서 유용)
                if len(doc) > 5:  # 페이지가 5개 이상인 경우에만
                    texts.append(f"--- 페이지 {page_num + 1} ---")
                texts.append("\n".join(page_text))
        
        logger.info(f"✅ PDF 파싱 완료: {os.path.basename(path)}")
        return "\n".join(texts)
    except Exception as e:
        logger.error(f"❌ PDF 파싱 오류 ({path}): {str(e)}")
        return f"PDF 파싱 오류: {path}\n오류 세부사항: {str(e)}"

def parse_docx(path: str) -> str:
    """DOCX 파일에서 텍스트 추출 - 스타일과 구조 정보 포함"""
    try:
        logger.info(f"📄 DOCX 파싱 시작: {os.path.basename(path)}")
        doc = docx.Document(path)
        full_text = []
        
        # 문서 속성 추출
        properties = doc.core_properties
        if properties.title:
            full_text.append(f"제목: {properties.title}")
        if properties.author:
            full_text.append(f"저자: {properties.author}")
        if properties.comments:
            full_text.append(f"설명: {properties.comments}")
        full_text.append("")  # 빈 줄 추가
        
        # 제목과 헤더 식별하여 구조적 텍스트 추출
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
                
            # 스타일 기반 구조 식별
            style = para.style.name
            text = para.text.strip()
            
            # 제목 스타일이면 마크다운 헤딩으로 변환
            if "Heading" in style:
                level = int(style.replace("Heading", "").strip()) if style.replace("Heading", "").strip().isdigit() else 1
                heading_marker = "#" * min(level, 6)
                full_text.append(f"{heading_marker} {text}")
            else:
                full_text.append(text)
        
        # 테이블 처리
        for i, table in enumerate(doc.tables):
            full_text.append(f"\n표 {i+1}:")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        logger.info(f"✅ DOCX 파싱 완료: {os.path.basename(path)}")
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"❌ DOCX 파싱 오류 ({path}): {str(e)}")
        return f"DOCX 파싱 오류: {path}\n오류 세부사항: {str(e)}"

def parse_txt(path: str) -> str:
    """TXT 파일에서 텍스트 추출 - 다양한 인코딩 지원"""
    logger.info(f"📄 TXT 파싱 시작: {os.path.basename(path)}")
    
    # 시도할 인코딩 목록
    encodings = ['utf-8', 'latin-1', 'cp949', 'euc-kr', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(path, 'r', encoding=encoding) as file:
                content = file.read()
            
            logger.info(f"✅ TXT 파싱 완료 ({encoding} 인코딩): {os.path.basename(path)}")
            
            # 파일 기본 정보 추가
            file_info = f"파일명: {os.path.basename(path)}\n"
            file_info += f"크기: {os.path.getsize(path) / 1024:.1f} KB\n\n"
            
            return file_info + content
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"❌ TXT 파싱 오류 ({path}, {encoding}): {str(e)}")
    
    # 모든 인코딩 시도 실패
    logger.error(f"❌ TXT 파싱 실패: 지원되는 인코딩을 찾을 수 없습니다 ({path})")
    return f"TXT 파싱 오류: {path}\n지원되는 인코딩을 찾을 수 없습니다."

def clean_text(text: str) -> str:
    """텍스트 정리 및 정규화 - 고급 정제 기능 추가"""
    if not text:
        return ""
        
    # 여러 줄바꿈 정리
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # 불필요한 공백 제거
    text = re.sub(r'\s{2,}', ' ', text)
    
    # 일반적인 쓸모없는 텍스트 제거 (예: 쿠키 정책, 구독 안내 등)
    patterns_to_remove = [
        r'쿠키를 사용.*?동의',
        r'Subscribe to.*?newsletter',
        r'구독.*?뉴스레터',
        r'Published:.*?\d{4}',
        r'Last modified on.*?\d{4}',
        r'Share on (?:Twitter|Facebook|LinkedIn)',
        r'\d+ shares',
        r'©.*?All rights reserved',
        r'Terms of (?:use|service)',
        r'Privacy Policy',
        r'All Rights Reserved',
        r'Please enable JavaScript',
        r'You need to enable JavaScript',
        r'ADVERTISEMENT',
        r'Advertisement',
        r'Sponsored Content',
        r'Click here to view'
    ]
    
    for pattern in patterns_to_remove:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # 중복된 텍스트 블록 제거 (특히 PDF에서 자주 발생)
    lines = text.split('\n')
    unique_lines = []
    seen_chunks = set()
    
    for line in lines:
        line = line.strip()
        if len(line) > 20:  # 긴 줄에 대해서만 중복 검사
            # 줄을 청크로 나누어 체크 (아주 긴 줄인 경우 일부만 중복될 수도 있음)
            chunk_size = 50
            chunks = [line[i:i+chunk_size] for i in range(0, len(line), chunk_size)]
            
            # 첫 번째 청크가 중복되면 건너뜀
            if chunks and chunks[0] in seen_chunks:
                continue
            
            # 청크 추가
            for chunk in chunks:
                if len(chunk) >= 20:  # 의미 있는 크기의 청크만 체크
                    seen_chunks.add(chunk)
        
        unique_lines.append(line)
    
    # 정리된 텍스트 반환
    cleaned_text = '\n'.join(unique_lines)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)  # 최종 줄바꿈 정리
    
    return cleaned_text.strip()

def parse_youtube_content(url: str) -> str:
    """YouTube 영상 콘텐츠 추출 - 동적 임포트 처리"""
    global youtube_parser
    
    # 필요시 유튜브 파서 모듈 임포트
    if youtube_parser is None:
        try:
            import youtube_parser as youtube_parser_module
            youtube_parser = youtube_parser_module
            logger.info("✅ YouTube 파서 모듈 로드 성공")
        except ImportError:
            logger.error("❌ YouTube 파서 모듈을 찾을 수 없습니다.")
            return f"YouTube 영상 파싱 오류 ({url}): YouTube 파서 모듈을 찾을 수 없습니다."
    
    try:
        return youtube_parser.parse_youtube(url)
    except Exception as e:
        logger.error(f"❌ YouTube 영상 파싱 오류 ({url}): {str(e)}")
        return f"YouTube 영상 파싱 오류 ({url}): {str(e)}"

# source_parser_updated.py 파일에 추가 또는 수정할 부분

def parse_sources(sources: List[SourceType], max_workers: int = 4) -> List[str]:
    """
    URL 또는 파일 목록 전체 처리 - 병렬 처리 및 오류 처리 강화
    
    Args:
        sources: URL 또는 파일 경로 목록
        max_workers: 병렬 처리 작업자 수
        
    Returns:
        파싱된 텍스트 목록
    """
    parsed_texts = []
    successful_sources = 0
    failed_sources = 0
    
    # 폴더 처리를 위한 소스 확장 (추가된 부분)
    expanded_sources = []
    for src in sources:
        # 이미지 폴더 처리
        if isinstance(src, dict) and src.get("type") == "image_folder":
            # 폴더 내 각 이미지 파일을 개별 소스로 추가
            for file_path in src.get("files", []):
                expanded_sources.append({
                    "type": file_path.split('.')[-1].lower(),
                    "path": file_path,
                    "ocr_engine": src.get("ocr_engine", "google")
                })
        else:
            # 일반 소스는 그대로 추가
            expanded_sources.append(src)
    
    # 확장된 소스 목록으로 교체
    sources = expanded_sources
    total = len(sources)
    
    logger.info(f"🔄 {total}개 소스 파싱 시작 (병렬 처리: {max_workers}개 작업자)")
    
    # 기존 코드는 그대로 유지...
    
    # 소스 타입에 따른 파싱 함수 매핑
    def parse_source(src_with_index: Tuple[int, SourceType]) -> Tuple[int, str, bool]:
        idx, src = src_with_index
        
        try:
            logger.info(f"[{idx+1}/{total}] 소스 파싱 중...")
            
            if isinstance(src, str):
                # URL 확인
                if src.startswith(('http://', 'https://')):
                    parsed = parse_url(src)
                else:
                    logger.warning(f"⚠️ 인식할 수 없는 소스 형식: {src}")
                    parsed = f"인식할 수 없는 소스 형식: {src}"
                    return idx, parsed, False
            elif isinstance(src, dict):
                src_type = src.get("type", "").lower()
                path = src.get("path", "")
                
                if not os.path.exists(path):
                    logger.error(f"❌ 파일을 찾을 수 없음: {path}")
                    return idx, f"파일을 찾을 수 없음: {path}", False
                
                if src_type == "pdf" or path.lower().endswith(".pdf"):
                    parsed = parse_pdf(path)
                elif src_type == "docx" or path.lower().endswith(".docx"):
                    parsed = parse_docx(path)
                elif src_type == "txt" or path.lower().endswith(".txt"):
                    parsed = parse_txt(path)
                elif any(path.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']):
                    # 기본 엔진으로 Google Vision 사용 (옵션으로 변경 가능)
                    engine = src.get("ocr_engine", "google")  # "google", "aws", "azure", "naver" 중 선택
                    
                    # OCR 엔진 검증
                    valid_engines = ["google", "aws", "azure", "naver"]
                    if engine not in valid_engines:
                        logger.warning(f"⚠️ 지원하지 않는 OCR 엔진: {engine}, 기본값 'google'로 설정합니다.")
                        engine = "google"
                    
                    logger.info(f"🔍 이미지 처리 중: {os.path.basename(path)} (OCR 엔진: {engine})")
                    parsed = parse_cloud_ocr(path, engine=engine)
                else:
                    logger.warning(f"⚠️ 지원하지 않는 파일 형식: {src}")
                    parsed = f"지원하지 않는 파일 형식: {src.get('path', '')}"
                    return idx, parsed, False
            else:
                logger.warning(f"⚠️ 알 수 없는 소스 유형: {type(src)}")
                parsed = ""
                return idx, parsed, False

            # 성공적으로 파싱된 텍스트 길이 확인
            success = parsed and len(parsed) > 100  # 최소 길이 기준
            return idx, parsed, success
            
        except Exception as e:
            logger.error(f"❌ 소스 파싱 중 예외 발생: {str(e)}")
            return idx, f"파싱 중 오류 발생: {str(e)}", False
    
    # 병렬 처리 실행
    results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(max_workers, total)) as executor:
        # 작업 제출
        future_to_idx = {executor.submit(parse_source, (i, src)): i for i, src in enumerate(sources)}
        
        # 결과 수집
        for future in concurrent.futures.as_completed(future_to_idx):
            try:
                idx, parsed_text, success = future.result()
                results.append((idx, parsed_text, success))
                
                if success:
                    successful_sources += 1
                    logger.info(f"✅ 소스 #{idx+1} 파싱 성공")
                else:
                    failed_sources += 1
                    logger.warning(f"⚠️ 소스 #{idx+1} 파싱 결과 불충분")
                
            except Exception as e:
                idx = future_to_idx[future]
                logger.error(f"❌ 소스 #{idx+1} 결과 처리 오류: {str(e)}")
                results.append((idx, f"파싱 중 예외 발생: {str(e)}", False))
                failed_sources += 1
    
    # 원래 순서대로 정렬
    results.sort(key=lambda x: x[0])
    parsed_texts = [result[1] for result in results]
    
    logger.info(f"🏁 소스 파싱 완료: 성공 {successful_sources}개, 실패 {failed_sources}개")
    
    return parsed_texts

if __name__ == "__main__":
    # 테스트 코드
    test_url = "https://en.wikipedia.org/wiki/Artificial_intelligence"
    logger.info("위키피디아 URL 테스트 중...")
    wiki_content = parse_url(test_url)
    logger.info(f"파싱된 내용 일부:\n{wiki_content[:500]}...")
    
    # YouTube 테스트
    youtube_url = "https://www.youtube.com/watch?v=dQw4w9WgXcQ"
    logger.info("\nYouTube URL 테스트 중...")
    try:
        yt_content = parse_url(youtube_url)
        logger.info(f"파싱된 내용 일부:\n{yt_content[:500]}...")
    except Exception as e:
        logger.error(f"YouTube 테스트 오류: {str(e)}")
    
    # 파일 파싱 테스트
    logger.info("\n다중 소스 병렬 파싱 테스트:")
    test_sources = [
        "https://en.wikipedia.org/wiki/Natural_language_processing",
        {"type": "txt", "path": "test_data/sample.txt"} if os.path.exists("test_data/sample.txt") else test_url
    ]
    
    parsed_results = parse_sources(test_sources)
    for i, result in enumerate(parsed_results):
        logger.info(f"소스 #{i+1} 결과 일부:\n{result[:200]}...")